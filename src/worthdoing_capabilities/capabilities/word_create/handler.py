"""Word document creation capability handler."""
from pathlib import Path


async def execute(input_data: dict) -> dict:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    output_path = Path(input_data["output_path"])
    content_blocks = input_data["content"]

    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Add title if provided
    title = input_data.get("title")
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks_count = 0
    char_count = 0

    for block in content_blocks:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            level = block.get("level", 1)
            text = block.get("text", "")
            doc.add_heading(text, level=min(level, 4))
            char_count += len(text)

        elif block_type == "paragraph":
            text = block.get("text", "")
            para = doc.add_paragraph(text)
            if block.get("bold"):
                for run in para.runs:
                    run.bold = True
            char_count += len(text)

        elif block_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Header row
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(h)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                # Data rows
                for row_idx, row_data in enumerate(rows, 1):
                    for col_idx, value in enumerate(row_data):
                        if col_idx < len(headers):
                            table.rows[row_idx].cells[col_idx].text = str(value)

        elif block_type == "list":
            items = block.get("items", [])
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
                char_count += len(item)

        blocks_count += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Rough page estimate: ~3000 chars per page
    page_estimate = max(1, char_count // 3000 + 1)

    return {
        "path": str(output_path),
        "blocks_count": blocks_count,
        "page_estimate": page_estimate,
    }
